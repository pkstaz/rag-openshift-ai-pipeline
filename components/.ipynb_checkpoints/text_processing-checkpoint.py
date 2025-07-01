"""
Text Processing Components para RAG Pipeline
Incluye: extract_text_component y chunk_text_component
"""

from kfp.dsl import component, Input, Output, Dataset

@component(
    base_image="pytorch/pytorch:2.0.1-cuda11.7-cudnn8-devel",
    packages_to_install=[
        "PyPDF2==3.0.1",
        "python-docx==0.8.11", 
        "minio==7.1.17",
        "chardet==5.2.0"
    ]
)
def extract_text_component(
    bucket_name: str,
    object_key: str,
    minio_endpoint: str,
    minio_access_key: str,
    minio_secret_key: str,
    extracted_text: Output[Dataset],
    metadata: Output[Dataset]
):
    """
    Extrae texto de documentos almacenados en MinIO.

    Args:
        bucket_name: Nombre del bucket en MinIO
        object_key: Path del archivo en el bucket
        minio_endpoint: Endpoint de MinIO
        minio_access_key: Access key de MinIO
        minio_secret_key: Secret key de MinIO
        extracted_text: Output dataset con el texto extraído
        metadata: Output dataset con metadata del documento
    """
    import os
    import json
    import tempfile
    from pathlib import Path
    from datetime import datetime
    from minio import Minio
    import PyPDF2
    from docx import Document
    import chardet

    # Conectar a MinIO
    minio_client = Minio(
        minio_endpoint,
        access_key=minio_access_key,
        secret_key=minio_secret_key,
        secure=False
    )

    # Crear directorio temporal
    with tempfile.TemporaryDirectory() as temp_dir:
        local_file_path = os.path.join(temp_dir, object_key.split('/')[-1])

        # Descargar archivo desde MinIO
        try:
            minio_client.fget_object(bucket_name, object_key, local_file_path)
            print(f"✅ Archivo descargado: {local_file_path}")
        except Exception as e:
            raise Exception(f"Error descargando archivo: {str(e)}")

        # Detectar tipo de archivo
        file_extension = Path(local_file_path).suffix.lower()
        file_size = os.path.getsize(local_file_path)

        # Extraer texto según el tipo de archivo
        extracted_content = ""

        if file_extension == '.pdf':
            try:
                with open(local_file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        extracted_content += page.extract_text() + "\n"
                print(f"✅ PDF procesado: {len(pdf_reader.pages)} páginas")
            except Exception as e:
                raise Exception(f"Error procesando PDF: {str(e)}")

        elif file_extension == '.docx':
            try:
                doc = Document(local_file_path)
                for paragraph in doc.paragraphs:
                    extracted_content += paragraph.text + "\n"
                print(f"✅ DOCX procesado: {len(doc.paragraphs)} párrafos")
            except Exception as e:
                raise Exception(f"Error procesando DOCX: {str(e)}")

        elif file_extension in ['.txt', '.md']:
            try:
                # Detectar encoding
                with open(local_file_path, 'rb') as file:
                    raw_data = file.read()
                    encoding = chardet.detect(raw_data)['encoding']

                # Leer con encoding detectado
                with open(local_file_path, 'r', encoding=encoding) as file:
                    extracted_content = file.read()
                print(f"✅ TXT procesado con encoding: {encoding}")
            except Exception as e:
                raise Exception(f"Error procesando TXT: {str(e)}")

        else:
            raise Exception(f"Tipo de archivo no soportado: {file_extension}")

        # Validar que se extrajo contenido
        if not extracted_content.strip():
            raise Exception("No se pudo extraer texto del documento")

        # Preparar metadata
        document_metadata = {
            "source_file": object_key,
            "file_type": file_extension,
            "file_size": file_size,
            "processed_at": datetime.now().isoformat(),
            "char_count": len(extracted_content),
            "word_count": len(extracted_content.split()),
            "bucket_name": bucket_name
        }

        # Guardar outputs
        with open(extracted_text.path, 'w', encoding='utf-8') as f:
            f.write(extracted_content)

        with open(metadata.path, 'w', encoding='utf-8') as f:
            json.dump(document_metadata, f, indent=2)

        print(f"✅ Texto extraído: {len(extracted_content)} caracteres")


@component(
    base_image="pytorch/pytorch:2.0.1-cuda11.7-cudnn8-devel",
    packages_to_install=[
        "tiktoken==0.5.1",
        "langchain==0.0.350"
    ]
)
def chunk_text_component(
    extracted_text: Input[Dataset],
    metadata: Input[Dataset],
    chunk_size: int,
    chunk_overlap: int,
    chunks: Output[Dataset]
):
    """
    Divide el texto en chunks con overlap para processing óptimo.

    Args:
        extracted_text: Input dataset con texto extraído
        metadata: Input dataset con metadata del documento
        chunk_size: Tamaño máximo de cada chunk (en tokens)
        chunk_overlap: Overlap entre chunks (en tokens)
        chunks: Output dataset con chunks procesados
    """
    import json
    import tiktoken
    from langchain.text_splitter import RecursiveCharacterTextSplitter

    # Leer input data
    with open(extracted_text.path, 'r', encoding='utf-8') as f:
        text_content = f.read()

    with open(metadata.path, 'r', encoding='utf-8') as f:
        doc_metadata = json.load(f)

    # Configurar tokenizer
    encoding = tiktoken.get_encoding("cl100k_base")

    def count_tokens(text: str) -> int:
        return len(encoding.encode(text))

    # Configurar text splitter
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size * 4,  # Aproximación: 1 token ≈ 4 caracteres
        chunk_overlap=chunk_overlap * 4,
        length_function=len,
        separators=["\n\n", "\n", ". ", " ", ""]
    )

    # Dividir texto en chunks
    text_chunks = text_splitter.split_text(text_content)
    print(f"✅ Texto dividido en {len(text_chunks)} chunks")

    # Procesar cada chunk
    processed_chunks = []

    for i, chunk_text in enumerate(text_chunks):
        token_count = count_tokens(chunk_text)

        chunk_metadata = {
            "chunk_id": f"{doc_metadata['source_file']}_chunk_{i:04d}",
            "chunk_index": i,
            "total_chunks": len(text_chunks),
            "text": chunk_text.strip(),
            "token_count": token_count,
            "char_count": len(chunk_text),
            "word_count": len(chunk_text.split()),
            "source_document": doc_metadata['source_file'],
            "file_type": doc_metadata['file_type'],
            "processed_at": doc_metadata['processed_at']
        }

        processed_chunks.append(chunk_metadata)

    # Filtrar chunks muy pequeños
    processed_chunks = [chunk for chunk in processed_chunks if chunk['token_count'] >= 10]

    print(f"✅ Chunks procesados: {len(processed_chunks)}")

    # Guardar chunks
    with open(chunks.path, 'w', encoding='utf-8') as f:
        json.dump(processed_chunks, f, indent=2, ensure_ascii=False)
